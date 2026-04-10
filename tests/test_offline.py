"""Offline test: parse -> format -> QC without LLM calls."""

import sys
import logging
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent.parent))
logging.basicConfig(level=logging.WARNING)

from agents.parser import DocumentParserAgent
from agents.formatter import DocumentFormatterAgent
from agents.quality import QualityControlAgent
from models.mapping import StructureMapping
from docx import Document

DRAFT = r"d:\desktop\Рабочие документы\СТО\Черновик стандарта организации.docx"
OUTPUT = r"d:\desktop\Рабочие документы\СТО\output\test_output.docx"


def main():
    parser = DocumentParserAgent()
    content = parser.run(draft_path=DRAFT)
    print(f"Parsed: {content.total_paragraphs} paras, {content.total_tables} tables, "
          f"{content.total_images} images, {len(content.sections)} sections")

    mapping = StructureMapping(
        sto_number="СТО 31025229–XXX–2025",
        document_title="Регламент сдачи ежемесячных материально-технических отчетов "
                       "материально ответственными лицами на основании выполненных работ",
        introduction_status="Введен впервые",
    )

    formatter = DocumentFormatterAgent()
    formatter.run(
        sections=content.sections,
        mapping=mapping,
        original=content,
        output_path=OUTPUT,
    )

    # Verify output
    doc = Document(OUTPUT)
    print(f"\nOutput: {len(doc.paragraphs)} paras, {len(doc.tables)} tables, {len(doc.sections)} sections")

    for i, sec in enumerate(doc.sections):
        w = round(sec.page_width / 914400 * 2.54, 1) if sec.page_width else 0
        h = round(sec.page_height / 914400 * 2.54, 1) if sec.page_height else 0
        l = round(sec.left_margin / 914400 * 2.54, 1) if sec.left_margin else 0
        r = round(sec.right_margin / 914400 * 2.54, 1) if sec.right_margin else 0
        hdr = ""
        if sec.header and sec.header.paragraphs:
            hdr = " | ".join(p.text.strip() for p in sec.header.paragraphs if p.text.strip())
        print(f"  Section {i}: {w}x{h}cm, L={l} R={r}, header='{hdr}'")

    # Font check
    fonts = set()
    sizes = set()
    for p in doc.paragraphs:
        for r in p.runs:
            if r.font.name:
                fonts.add(r.font.name)
            if r.font.size:
                sizes.add(round(r.font.size / 12700, 1))
    print(f"\nFonts used: {fonts}")
    print(f"Sizes used: {sizes}")

    # Title page check
    first_texts = [p.text.strip() for p in doc.paragraphs[:20] if p.text.strip()]
    print(f"\nTitle page first lines: {first_texts[:8]}")

    # QC
    qc = QualityControlAgent(llm=None)
    report = qc.run(output_path=OUTPUT)
    print(f"\n{report.summary()}")


if __name__ == "__main__":
    main()
