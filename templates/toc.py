"""Generate a Table of Contents page."""

from __future__ import annotations

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from models.document import SectionData
from templates.formatting_rules import FONT_NAME, FONT_SIZE_BODY, FONT_SIZE_TOC_TITLE


def add_toc_page(
    doc: Document,
    sections: list[SectionData],
    *,
    appendix_titles: list[str] | None = None,
) -> None:
    """Add a 'Содержание' page with section listings.

    Since python-docx cannot compute page numbers, we insert a TOC field
    that Word will update on open, plus a static fallback listing.
    """
    doc.add_page_break()

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    r = p.add_run("Содержание")
    r.font.name = FONT_NAME
    r.font.size = FONT_SIZE_TOC_TITLE
    r.bold = True

    # Insert a Word TOC field (updated when document is opened in Word)
    _insert_toc_field(doc)

    # Static fallback listing
    for sec in sections:
        _add_toc_entry(doc, sec.number, sec.title, level=1)
        for sub in sec.subsections:
            _add_toc_entry(doc, sub.number, sub.title, level=2)

    if appendix_titles:
        for title in appendix_titles:
            _add_toc_entry(doc, "", title, level=1)

    # Standard mandatory trailing sheets
    _add_toc_entry(doc, "", "Лист регистрации изменений", level=1)
    _add_toc_entry(doc, "", "Лист ознакомления", level=1)
    _add_toc_entry(doc, "", "Лист согласования", level=1)


def _add_toc_entry(doc: Document, number: str, title: str, *, level: int = 1) -> None:
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(2)

    indent_chars = "    " * (level - 1)
    prefix = f"{number} " if number else ""
    text = f"{indent_chars}{prefix}{title}"

    r = p.add_run(text)
    r.font.name = FONT_NAME
    r.font.size = FONT_SIZE_BODY


def _insert_toc_field(doc: Document) -> None:
    """Insert an auto-updating TOC field (Word will refresh on open)."""
    p = doc.add_paragraph()
    run = p.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    run._element.append(fld_char_begin)

    run2 = p.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    run2._element.append(instr)

    run3 = p.add_run()
    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    run3._element.append(fld_char_separate)

    run4 = p.add_run()
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run4._element.append(fld_char_end)
