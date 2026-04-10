"""Generate the title page (титульный лист) per СТО 31025229-001-2024 Appendix В."""

from __future__ import annotations

from datetime import datetime
import textwrap

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from config import COMPANY_NAME, DIRECTOR_NAME, DIRECTOR_TITLE
from templates.formatting_rules import (
    FONT_NAME,
    FONT_SIZE_TITLE,
    STO_NUMBER_INDENT,
)


def add_title_page(
    doc: Document,
    *,
    sto_number: str,
    document_title: str,
    intro_status: str = "Введен впервые",
    year: int | None = None,
) -> None:
    """Add a complete title page to the document."""
    if year is None:
        year = datetime.now().year

    _add_approve_block(doc, year)
    _add_empty(doc, count=3)
    _add_sto_number(doc, sto_number)
    _add_empty(doc, count=3)
    _add_centered_bold(doc, "СТАНДАРТ ОРГАНИЗАЦИИ")
    _add_empty(doc)
    _add_line(doc)
    _add_empty(doc, count=2)
    _add_left_bold(doc, "Система менеджмента качества")
    _add_empty(doc)
    _add_document_title(doc, document_title, intro_status)
    _add_empty(doc)
    _add_line(doc)
    _add_empty(doc, count=2)
    _add_intro_block(doc)
    # Keep company/year near the bottom as in STO sample.
    _add_empty(doc, count=12)
    _add_centered(doc, COMPANY_NAME)
    _add_empty(doc, count=1)
    _add_centered(doc, str(year))


def _run(para, text: str, *, bold: bool = False, size=FONT_SIZE_TITLE):
    r = para.add_run(text)
    r.font.name = FONT_NAME
    r.font.size = size
    r.bold = bold
    return r


def _add_approve_block(doc: Document, year: int) -> None:
    lines = [
        "УТВЕРЖДАЮ",
        DIRECTOR_TITLE,
        COMPANY_NAME,
        f"______________ {DIRECTOR_NAME}",
        f"«_____»____________{year} г.",
    ]
    for line in lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        _run(p, line, bold=True)


def _add_sto_number(doc: Document, sto_number: str) -> None:
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.first_line_indent = STO_NUMBER_INDENT
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    normalized = sto_number.replace("СТО ", "").strip()
    _run(p, f"СТО {normalized}", bold=True)


def _add_centered_bold(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    _run(p, text, bold=True)


def _add_centered(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    _run(p, text, bold=False)


def _add_left_bold(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    _run(p, text, bold=True)


def _add_line(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    _run(p, "─" * 72, bold=True)


def _add_document_title(doc: Document, title: str, intro_status: str) -> None:
    title_upper = " ".join(title.upper().split())
    lines = textwrap.wrap(title_upper, width=45)
    if not lines:
        lines = [title_upper]
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    _remove_table_borders(table)
    left, right = table.rows[0].cells
    left.width = Cm(13.0)
    right.width = Cm(5.0)

    left_p = left.paragraphs[0]
    left_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    left_pf = left_p.paragraph_format
    left_pf.space_before = Pt(0)
    left_pf.space_after = Pt(0)
    for i, line in enumerate(lines):
        p = left_p if i == 0 else left.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        _run(p, line, bold=True)

    right_p = right.paragraphs[0]
    right_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right_pf = right_p.paragraph_format
    right_pf.space_before = Pt(0)
    right_pf.space_after = Pt(0)
    _run(right_p, intro_status, bold=True)


def _add_intro_block(doc: Document) -> None:
    lines = [
        "Введен в действие Приказом",
        "от _____________№ ______",
    ]
    for line in lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf = p.paragraph_format
        pf.first_line_indent = Cm(1.0)
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        _run(p, line)

    _add_empty(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    _run(p, "Дата введения ____________")


def _add_empty(doc: Document, count: int = 1) -> None:
    for _ in range(count):
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        r = p.add_run("")
        r.font.name = FONT_NAME
        r.font.size = FONT_SIZE_TITLE


def _remove_table_borders(table) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)

    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        borders.append(el)
    tbl_pr.append(borders)
