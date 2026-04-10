"""Generate mandatory appendix sheets per СТО 31025229-001-2024."""

from __future__ import annotations

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from templates.formatting_rules import FONT_NAME, FONT_SIZE_BODY, FONT_SIZE_HEADING


def add_change_registration_sheet(doc: Document) -> None:
    """Лист регистрации изменений (Appendix Д of the standard)."""
    doc.add_page_break()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(12)
    r = p.add_run("Лист регистрации изменений")
    r.font.name = FONT_NAME
    r.font.size = FONT_SIZE_HEADING
    r.bold = True

    headers = [
        "Изм.",
        "Номера листов (страниц)\nизмененных",
        "Номера листов (страниц)\nзамененных",
        "Номера листов (страниц)\nновых",
        "Номера листов (страниц)\nаннулированных",
        "Всего\nлистов\n(страниц)\nв докум.",
        "№\nдокумента",
        "Подпись",
        "Дата",
    ]

    num_rows = 6
    table = doc.add_table(rows=num_rows, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _apply_table_style(table)

    for ci, hdr in enumerate(headers):
        cell = table.cell(0, ci)
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(hdr)
        r.font.name = FONT_NAME
        r.font.size = Pt(10)
        r.bold = True

    for ri in range(1, num_rows):
        for ci in range(len(headers)):
            cell = table.cell(ri, ci)
            cell.text = ""
            p = cell.paragraphs[0]
            r = p.add_run("")
            r.font.name = FONT_NAME
            r.font.size = Pt(10)


def add_familiarization_sheet(doc: Document) -> None:
    """Лист ознакомления (Appendix Е of the standard)."""
    doc.add_page_break()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(12)
    r = p.add_run("Лист ознакомления")
    r.font.name = FONT_NAME
    r.font.size = FONT_SIZE_HEADING
    r.bold = True

    headers = ["№п/п", "Ф.И.О.", "Должность", "Подпись", "Дата"]
    num_rows = 22

    table = doc.add_table(rows=num_rows, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _apply_table_style(table)

    for ci, hdr in enumerate(headers):
        cell = table.cell(0, ci)
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(hdr)
        r.font.name = FONT_NAME
        r.font.size = FONT_SIZE_BODY
        r.bold = True

    for ri in range(1, num_rows):
        cell = table.cell(ri, 0)
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(str(ri))
        r.font.name = FONT_NAME
        r.font.size = FONT_SIZE_BODY

        for ci in range(1, len(headers)):
            cell = table.cell(ri, ci)
            cell.text = ""
            p = cell.paragraphs[0]
            r = p.add_run("")
            r.font.name = FONT_NAME
            r.font.size = FONT_SIZE_BODY


def add_approval_sheet(doc: Document, signatories: list[tuple[str, str]] | None = None) -> None:
    """Лист согласования."""
    doc.add_page_break()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(12)
    r = p.add_run("Лист согласования")
    r.font.name = FONT_NAME
    r.font.size = FONT_SIZE_HEADING
    r.bold = True

    if signatories is None:
        signatories = [
            ("Ответственный разработчик:", ""),
            ("ПРП по СМК", ""),
        ]

    for title, name in signatories:
        _add_empty(doc)

        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        r = p.add_run(title)
        r.font.name = FONT_NAME
        r.font.size = FONT_SIZE_BODY

        if name:
            p = doc.add_paragraph()
            pf = p.paragraph_format
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
            r = p.add_run(f"______________ {name}")
            r.font.name = FONT_NAME
            r.font.size = FONT_SIZE_BODY

        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        r = p.add_run("«_____»  __________ 20____г.")
        r.font.name = FONT_NAME
        r.font.size = FONT_SIZE_BODY


def _apply_table_style(table) -> None:
    """Apply borders to all cells."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn as _qn

    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")

    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(_qn("w:val"), "single")
        el.set(_qn("w:sz"), "4")
        el.set(_qn("w:space"), "0")
        el.set(_qn("w:color"), "000000")
        borders.append(el)
    tbl_pr.append(borders)


def _add_empty(doc: Document) -> None:
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    r = p.add_run("")
    r.font.name = FONT_NAME
    r.font.size = FONT_SIZE_BODY
