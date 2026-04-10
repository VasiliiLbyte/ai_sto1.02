"""Agent 4: DocumentFormatter — assemble the final DOCX document."""

from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from agents.base import BaseAgent
from models.document import DocumentContent, SectionData
from models.mapping import StructureMapping

_SECTION_NUM_RE = re.compile(r"^(\d+(?:\.\d+)*)\.*\s+", re.UNICODE)
from services.docx_writer import DocxWriter
from templates.appendix_sheets import (
    add_approval_sheet,
    add_change_registration_sheet,
    add_familiarization_sheet,
)
from templates.title_page import add_title_page
from templates.toc import add_toc_page


class DocumentFormatterAgent(BaseAgent):
    name = "DocumentFormatter"

    def run(
        self,
        *,
        sections: list[SectionData],
        mapping: StructureMapping,
        original: DocumentContent,
        output_path: str,
        **kwargs: Any,
    ) -> Path:
        self.logger.info("Assembling formatted document …")
        writer = DocxWriter()

        sto_number = mapping.sto_number or original.sto_number or "СТО 31025229–XXX–2025"
        doc_title = mapping.document_title or original.title or "Стандарт организации"
        intro_status = mapping.introduction_status or "Введен впервые"

        # 1. Title page
        add_title_page(
            writer.doc,
            sto_number=sto_number,
            document_title=doc_title,
            intro_status=intro_status,
        )

        # 2. Table of contents
        appendix_titles = self._collect_appendix_titles(original)
        add_toc_page(writer.doc, sections, appendix_titles=appendix_titles)

        # 3. Header & footer (on a new section so title page has no header)
        self._setup_body_section(writer, sto_number)

        # 4. Body sections
        for sec in sections:
            writer.write_section(sec)

        # 5. Appendices from the original document
        self._write_appendices(writer, original)

        # 6. Mandatory sheets
        add_change_registration_sheet(writer.doc)
        add_familiarization_sheet(writer.doc)
        add_approval_sheet(writer.doc)

        saved = writer.save(output_path)
        self.logger.info("Document saved: %s", saved)
        return saved

    @staticmethod
    def _setup_body_section(writer: DocxWriter, sto_number: str) -> None:
        """Add a new section for the body so the title page stays header-free."""
        from docx.enum.section import WD_ORIENT
        from templates.formatting_rules import (
            MARGIN_BOTTOM,
            MARGIN_LEFT,
            MARGIN_RIGHT,
            MARGIN_TOP,
            PAGE_HEIGHT,
            PAGE_WIDTH,
        )

        writer.doc.add_section()
        new_sec = writer.doc.sections[-1]
        new_sec.page_width = PAGE_WIDTH
        new_sec.page_height = PAGE_HEIGHT
        new_sec.left_margin = MARGIN_LEFT
        new_sec.right_margin = MARGIN_RIGHT
        new_sec.top_margin = MARGIN_TOP
        new_sec.bottom_margin = MARGIN_BOTTOM
        new_sec.orientation = WD_ORIENT.PORTRAIT

        # Unlink header/footer from previous (title page) section
        new_sec.header.is_linked_to_previous = False
        new_sec.footer.is_linked_to_previous = False

        writer.setup_header_footer(sto_number)

    @staticmethod
    def _collect_appendix_titles(original: DocumentContent) -> list[str]:
        start = original.appendix_start_index
        if start < 0:
            start = 0
        seen: set[str] = set()
        titles: list[str] = []
        for p in original.paragraphs:
            if p.index < start:
                continue
            t = p.text.strip()
            if not t.lower().startswith("приложение"):
                continue
            if t.lower().strip(".: ") == "приложения":
                continue
            key = t.lower().replace(" ", "").replace(".", "").replace("№", "")
            if key in seen:
                continue
            seen.add(key)
            titles.append(t)
        return titles

    def _write_appendices(self, writer: DocxWriter, original: DocumentContent) -> None:
        """Copy appendix content starting from appendix_start_index.

        Uses the paragraph index boundary detected by the parser to avoid
        accidentally capturing TOC or body content. Interleaves tables and
        images at their original positions.
        """
        start = original.appendix_start_index
        if start < 0:
            self.logger.warning("No appendix boundary detected; skipping appendix writing")
            return

        # Build lookup of appendix tables/images by position
        placed_in_body = self._body_element_indices(original)
        appendix_tables: dict[int, list] = {}
        for tbl in original.tables:
            if tbl.index not in placed_in_body["tables"]:
                appendix_tables.setdefault(tbl.position_after_paragraph, []).append(tbl)
        appendix_images: dict[int, list] = {}
        for img in original.images:
            if img.index not in placed_in_body["images"]:
                appendix_images.setdefault(img.position_after_paragraph, []).append(img)

        capturing = False
        seen_titles: set[str] = set()
        placed_img_indices: set[int] = set()
        placed_tbl_indices: set[int] = set()

        for p in original.paragraphs:
            if p.index < start:
                continue

            t = p.text.strip()
            has_bold = any(r.bold for r in p.runs) if p.runs else False

            if t.startswith("*М-") or t.startswith("*M-") or t.startswith("*Справочно"):
                break
            if t.startswith("*Примечание"):
                break

            # Skip body section headings that appear between appendix blocks
            if has_bold and _SECTION_NUM_RE.match(t):
                capturing = False
                continue

            if t.lower().startswith("приложение"):
                if t.lower().strip(".: ") == "приложения":
                    continue
                key = t.lower().replace(" ", "").replace(".", "").replace("№", "")
                if key in seen_titles:
                    continue
                seen_titles.add(key)
                capturing = True
                writer.add_page_break()
                writer.add_section_heading("", t)

                self._insert_elements_at(
                    writer, p.index, appendix_tables, appendix_images,
                    placed_tbl_indices, placed_img_indices,
                )
                continue

            if capturing and t:
                writer.add_paragraph_data(p)

            if capturing:
                self._insert_elements_at(
                    writer, p.index, appendix_tables, appendix_images,
                    placed_tbl_indices, placed_img_indices,
                )

        # Any remaining unplaced appendix elements
        for pos in sorted(appendix_tables.keys()):
            for tbl in appendix_tables[pos]:
                if tbl.index not in placed_tbl_indices:
                    writer.add_table(tbl)
                    placed_tbl_indices.add(tbl.index)
        for pos in sorted(appendix_images.keys()):
            for img in appendix_images[pos]:
                if img.index not in placed_img_indices:
                    writer.add_image(img)
                    placed_img_indices.add(img.index)

    @staticmethod
    def _insert_elements_at(
        writer: DocxWriter,
        para_index: int,
        tables: dict[int, list],
        images: dict[int, list],
        placed_tbls: set[int],
        placed_imgs: set[int],
    ) -> None:
        for tbl in tables.get(para_index, []):
            if tbl.index not in placed_tbls:
                writer.add_table(tbl)
                placed_tbls.add(tbl.index)
        for img in images.get(para_index, []):
            if img.index not in placed_imgs:
                writer.add_image(img)
                placed_imgs.add(img.index)

    @staticmethod
    def _body_element_indices(original: DocumentContent) -> dict[str, set[int]]:
        """Collect indices of tables/images already placed in body sections."""
        tables: set[int] = set()
        images: set[int] = set()
        for sec in original.sections:
            for tbl in sec.tables:
                tables.add(tbl.index)
            for img in sec.images:
                images.add(img.index)
            for sub in sec.subsections:
                for tbl in sub.tables:
                    tables.add(tbl.index)
                for img in sub.images:
                    images.add(img.index)
        return {"tables": tables, "images": images}
