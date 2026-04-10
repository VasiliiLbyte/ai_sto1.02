"""Agent 5: QualityController — validate the output document against STO rules."""

from __future__ import annotations

import logging
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Pt

from agents.base import BaseAgent
from config import MODEL_QC
from services.llm_client import LLMClient

logger = logging.getLogger(__name__)

EXPECTED_FONT = "Times New Roman"
EXPECTED_BODY_SIZE_PT = 13.0
EXPECTED_HEADING_SIZE_PT = 14.0
EXPECTED_INDENT_CM = 1.0
EXPECTED_MARGINS_CM = {"left": 2.0, "right": 1.0, "top": 2.0, "bottom": 2.0}

REQUIRED_SECTIONS_KEYWORDS = [
    "область применения",
    "нормативные ссылки",
    "термины",
    "ответственность",
]

REQUIRED_SHEETS = [
    "лист регистрации изменений",
    "лист ознакомления",
    "лист согласования",
]


class QualityIssue:
    def __init__(self, severity: str, message: str):
        self.severity = severity  # "error" | "warning"
        self.message = message

    def __repr__(self) -> str:
        return f"[{self.severity.upper()}] {self.message}"


class QualityReport:
    def __init__(self) -> None:
        self.issues: list[QualityIssue] = []

    @property
    def passed(self) -> bool:
        return not any(i.severity == "error" for i in self.issues)

    def add_error(self, msg: str) -> None:
        self.issues.append(QualityIssue("error", msg))

    def add_warning(self, msg: str) -> None:
        self.issues.append(QualityIssue("warning", msg))

    def summary(self) -> str:
        errors = [i for i in self.issues if i.severity == "error"]
        warnings = [i for i in self.issues if i.severity == "warning"]
        lines = [
            f"Quality Report: {len(errors)} errors, {len(warnings)} warnings",
            f"Status: {'PASS' if self.passed else 'FAIL'}",
            "",
        ]
        for issue in self.issues:
            lines.append(str(issue))
        return "\n".join(lines)


class QualityControlAgent(BaseAgent):
    name = "QualityController"

    def __init__(self, llm: LLMClient | None = None) -> None:
        super().__init__()
        self.llm = llm

    def run(self, *, output_path: str, original_text: str = "", **kwargs: Any) -> QualityReport:
        self.logger.info("Validating output document: %s", output_path)
        report = QualityReport()

        doc = Document(output_path)

        self._check_page_setup(doc, report)
        self._check_fonts(doc, report)
        self._check_structure(doc, report)
        self._check_headers_footers(doc, report)
        self._check_mandatory_sheets(doc, report)

        if self.llm and original_text:
            self._check_content_preservation(doc, original_text, report)

        self.logger.info("Validation complete: %s", "PASS" if report.passed else "FAIL")
        return report

    def _check_page_setup(self, doc: Document, report: QualityReport) -> None:
        for i, sec in enumerate(doc.sections):
            if sec.page_width:
                w_cm = round(sec.page_width / 914400 * 2.54, 1)
                h_cm = round(sec.page_height / 914400 * 2.54, 1)
                if abs(w_cm - 21.0) > 0.5 and abs(h_cm - 21.0) > 0.5:
                    report.add_error(f"Section {i}: page size {w_cm}x{h_cm}cm, expected 21.0x29.7 or 29.7x21.0")

            if sec.left_margin:
                margins = {
                    "left": round(sec.left_margin / 914400 * 2.54, 1),
                    "right": round(sec.right_margin / 914400 * 2.54, 1),
                    "top": round(sec.top_margin / 914400 * 2.54, 1),
                    "bottom": round(sec.bottom_margin / 914400 * 2.54, 1),
                }
                for side, val in margins.items():
                    expected = EXPECTED_MARGINS_CM[side]
                    if abs(val - expected) > 0.3:
                        report.add_warning(
                            f"Section {i}: {side} margin = {val}cm, expected {expected}cm"
                        )

    def _check_fonts(self, doc: Document, report: QualityReport) -> None:
        wrong_fonts: dict[str, int] = {}
        wrong_sizes: dict[float, int] = {}
        total_runs = 0

        for p in doc.paragraphs:
            for r in p.runs:
                total_runs += 1
                if r.font.name and r.font.name != EXPECTED_FONT:
                    wrong_fonts[r.font.name] = wrong_fonts.get(r.font.name, 0) + 1
                if r.font.size:
                    size_pt = round(r.font.size / 12700, 1)
                    if size_pt not in (EXPECTED_BODY_SIZE_PT, EXPECTED_HEADING_SIZE_PT, 10.0, 8.0):
                        wrong_sizes[size_pt] = wrong_sizes.get(size_pt, 0) + 1

        for font_name, count in wrong_fonts.items():
            pct = count / max(total_runs, 1) * 100
            if pct > 5:
                report.add_error(f"Font '{font_name}' used in {count} runs ({pct:.0f}%) — expected {EXPECTED_FONT}")
            else:
                report.add_warning(f"Font '{font_name}' in {count} runs")

        for size, count in wrong_sizes.items():
            if count > 5:
                report.add_warning(f"Non-standard font size {size}pt in {count} runs")

    def _check_structure(self, doc: Document, report: QualityReport) -> None:
        all_text = "\n".join(p.text.lower() for p in doc.paragraphs)
        for kw in REQUIRED_SECTIONS_KEYWORDS:
            if kw not in all_text:
                report.add_error(f"Missing required section containing '{kw}'")

    def _check_headers_footers(self, doc: Document, report: QualityReport) -> None:
        has_header = False
        has_footer = False
        for sec in doc.sections:
            if sec.header and sec.header.paragraphs:
                for p in sec.header.paragraphs:
                    if p.text.strip():
                        has_header = True
            if sec.footer and sec.footer.paragraphs:
                for p in sec.footer.paragraphs:
                    if p.text.strip() or p._element.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fldChar"):
                        has_footer = True

        if not has_header:
            report.add_warning("No header with STO number found")
        if not has_footer:
            report.add_warning("No footer with page number found")

    def _check_mandatory_sheets(self, doc: Document, report: QualityReport) -> None:
        all_text = "\n".join(p.text.lower() for p in doc.paragraphs)
        for sheet in REQUIRED_SHEETS:
            if sheet not in all_text:
                report.add_error(f"Missing mandatory sheet: '{sheet}'")

    def _check_content_preservation(
        self, doc: Document, original_text: str, report: QualityReport
    ) -> None:
        """Use LLM to verify that the meaning was preserved."""
        if not self.llm:
            return

        output_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())[:8000]
        original_snippet = original_text[:8000]

        prompt = (
            "Сравни два текста: исходный черновик и итоговый документ.\n"
            "Проверь, сохранена ли смысловая нагрузка. "
            "Ответь JSON: {\"preserved\": true/false, \"issues\": [\"...\"]}\n\n"
            f"ИСХОДНЫЙ ТЕКСТ (фрагмент):\n{original_snippet}\n\n"
            f"ИТОГОВЫЙ ТЕКСТ (фрагмент):\n{output_text}"
        )

        try:
            data = self.llm.chat_json(
                model=MODEL_QC,
                messages=[
                    {"role": "system", "content": "Ты — контролер качества документов СТО."},
                    {"role": "user", "content": prompt},
                ],
            )
            if not data.get("preserved", True):
                for issue in data.get("issues", []):
                    report.add_error(f"Content issue: {issue}")
        except Exception as e:
            report.add_warning(f"LLM content check failed: {e}")
