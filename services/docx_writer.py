"""Build a DOCX file with exact STO 31025229-001-2024 formatting."""

from __future__ import annotations

import logging
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from models.document import (
    ImageData,
    ParagraphData,
    SectionData,
    TableCellData,
    TableData,
)
from templates.formatting_rules import (
    FIRST_LINE_INDENT,
    FONT_NAME,
    FONT_SIZE_BODY,
    FONT_SIZE_HEADING,
    HEADING_SPACE_AFTER,
    LINE_SPACING_SINGLE,
    MARGIN_BOTTOM,
    MARGIN_LEFT,
    MARGIN_RIGHT,
    MARGIN_TOP,
    PAGE_HEIGHT,
    PAGE_WIDTH,
)

logger = logging.getLogger(__name__)

_ALIGN_MAP = {
    "LEFT": WD_ALIGN_PARAGRAPH.LEFT,
    "CENTER": WD_ALIGN_PARAGRAPH.CENTER,
    "RIGHT": WD_ALIGN_PARAGRAPH.RIGHT,
    "JUSTIFY": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


class DocxWriter:
    """Builds a formatted DOCX document."""

    def __init__(self) -> None:
        self.doc = Document()
        self._setup_default_style()
        self._setup_page()

    # ------------------------------------------------------------------
    # Page & style setup
    # ------------------------------------------------------------------

    def _setup_page(self) -> None:
        section = self.doc.sections[0]
        section.page_width = PAGE_WIDTH
        section.page_height = PAGE_HEIGHT
        section.left_margin = MARGIN_LEFT
        section.right_margin = MARGIN_RIGHT
        section.top_margin = MARGIN_TOP
        section.bottom_margin = MARGIN_BOTTOM
        section.orientation = WD_ORIENT.PORTRAIT

    def _setup_default_style(self) -> None:
        style = self.doc.styles["Normal"]
        font = style.font
        font.name = FONT_NAME
        font.size = FONT_SIZE_BODY
        pf = style.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing = LINE_SPACING_SINGLE

    # ------------------------------------------------------------------
    # Headers & footers
    # ------------------------------------------------------------------

    def setup_header_footer(self, sto_number: str, section_index: int = -1) -> None:
        """Add STO number in header-right, page number in footer-right.

        By default operates on the last section so the title-page section
        (index 0) stays header-free.
        """
        section = self.doc.sections[section_index]

        # Header
        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = hp.add_run(sto_number)
        r.font.name = FONT_NAME
        r.font.size = FONT_SIZE_BODY

        # Footer with auto page number
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _add_page_number_field(fp)

    # ------------------------------------------------------------------
    # Content writing
    # ------------------------------------------------------------------

    def add_section_heading(self, number: str, title: str) -> None:
        """Add a section heading: '5 Заголовок' — 14pt bold, JUSTIFY."""
        text = f"{number} {title}" if number else title
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf = p.paragraph_format
        pf.first_line_indent = FIRST_LINE_INDENT
        pf.space_before = Pt(0)
        pf.space_after = HEADING_SPACE_AFTER
        r = p.add_run(text)
        r.font.name = FONT_NAME
        r.font.size = FONT_SIZE_HEADING
        r.bold = True

    def add_body_paragraph(
        self,
        text: str,
        *,
        bold: bool = False,
        italic: bool = False,
        indent: bool = True,
        alignment: str = "JUSTIFY",
    ) -> None:
        """Add a standard body paragraph — 13pt, JUSTIFY, indent 1.0cm."""
        p = self.doc.add_paragraph()
        p.alignment = _ALIGN_MAP.get(alignment, WD_ALIGN_PARAGRAPH.JUSTIFY)
        pf = p.paragraph_format
        if indent:
            pf.first_line_indent = FIRST_LINE_INDENT
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        r = p.add_run(text)
        r.font.name = FONT_NAME
        r.font.size = FONT_SIZE_BODY
        r.bold = bold
        r.italic = italic

    def add_term_paragraph(self, term: str, definition: str) -> None:
        """Add a term definition: **term**: definition."""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf = p.paragraph_format
        pf.first_line_indent = FIRST_LINE_INDENT
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)

        r1 = p.add_run(f"{term}: ")
        r1.font.name = FONT_NAME
        r1.font.size = FONT_SIZE_BODY
        r1.bold = True

        r2 = p.add_run(definition)
        r2.font.name = FONT_NAME
        r2.font.size = FONT_SIZE_BODY

    def add_paragraph_data(self, para: ParagraphData) -> None:
        """Write a ParagraphData with its original run formatting adjusted to STO."""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf = p.paragraph_format
        pf.first_line_indent = FIRST_LINE_INDENT
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)

        if para.runs:
            for rd in para.runs:
                r = p.add_run(rd.text)
                r.font.name = FONT_NAME
                r.font.size = FONT_SIZE_BODY
                r.bold = rd.bold
                r.italic = rd.italic
        else:
            r = p.add_run(para.text)
            r.font.name = FONT_NAME
            r.font.size = FONT_SIZE_BODY

    def add_table(self, table: TableData) -> None:
        """Copy a table from parsed data with STO formatting."""
        if not table.rows:
            return

        num_rows = len(table.rows)
        num_cols = max(len(row) for row in table.rows) if table.rows else 0
        if num_cols == 0:
            return

        tbl = self.doc.add_table(rows=num_rows, cols=num_cols)
        _apply_borders(tbl)

        for ri, row in enumerate(table.rows):
            for ci, cell_data in enumerate(row):
                if ci >= num_cols:
                    break
                cell = tbl.cell(ri, ci)
                cell.text = ""
                for pi, cp in enumerate(cell_data.paragraphs):
                    if pi == 0:
                        p = cell.paragraphs[0]
                    else:
                        p = cell.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for rd in cp.runs:
                        r = p.add_run(rd.text)
                        r.font.name = FONT_NAME
                        r.font.size = FONT_SIZE_BODY
                        r.bold = rd.bold
                        r.italic = rd.italic
                    if not cp.runs:
                        r = p.add_run(cp.text)
                        r.font.name = FONT_NAME
                        r.font.size = FONT_SIZE_BODY

    def add_image(self, image: ImageData) -> None:
        """Insert an image preserving its original dimensions."""
        img_bytes = image.get_bytes()
        if not img_bytes:
            return

        # Max usable width = page_width - left_margin - right_margin = 18 cm
        max_width_cm = 18.0
        w_cm = image.width_cm if image.width_cm else 15.0
        if w_cm > max_width_cm:
            w_cm = max_width_cm

        # EMF/WMF images may not be directly insertable; try conversion
        bio = BytesIO(img_bytes)
        if image.content_type in ("image/x-emf", "image/x-wmf"):
            bio = _try_convert_metafile(img_bytes) or bio

        try:
            self.doc.add_picture(bio, width=Cm(w_cm))
            last_p = self.doc.paragraphs[-1]
            last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            logger.warning("Failed to insert image %s: %s", image.filename, e)

    def add_page_break(self) -> None:
        self.doc.add_page_break()

    # ------------------------------------------------------------------
    # Section writing (high-level)
    # ------------------------------------------------------------------

    def write_section(self, section: SectionData) -> None:
        """Write a full section with heading, paragraphs, tables, and images."""
        self.add_section_heading(section.number, section.title)

        img_positions: dict[int, ImageData] = {}
        for img in section.images:
            img_positions[img.position_after_paragraph] = img

        tbl_positions: dict[int, TableData] = {}
        for tbl in section.tables:
            tbl_positions[tbl.position_after_paragraph] = tbl

        for para in section.paragraphs:
            if para.text.strip():
                self.add_paragraph_data(para)

            if para.index in img_positions:
                self.add_image(img_positions[para.index])
            if para.index in tbl_positions:
                self.add_table(tbl_positions[para.index])

        # Tables/images not positioned after any paragraph
        for tbl in section.tables:
            if tbl.position_after_paragraph not in {p.index for p in section.paragraphs}:
                self.add_table(tbl)
        for img in section.images:
            if img.position_after_paragraph not in {p.index for p in section.paragraphs}:
                self.add_image(img)

        for sub in section.subsections:
            self.write_section(sub)

    # ------------------------------------------------------------------
    # Save
    # ------------------------------------------------------------------

    def save(self, path: str | Path) -> Path:
        path = Path(path)
        path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(path))
        logger.info("Document saved to %s", path)
        return path


# ======================================================================
# Helpers
# ======================================================================

def _add_page_number_field(paragraph) -> None:
    """Insert a PAGE field into a paragraph."""
    run = paragraph.add_run()
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE_BODY

    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    run._element.append(fld_char_begin)

    run2 = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run2._element.append(instr)

    run3 = paragraph.add_run()
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run3._element.append(fld_char_end)


def _apply_borders(table) -> None:
    """Apply single borders to all cells of a table."""
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)

    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
        borders.append(el)
    tbl_pr.append(borders)


def _try_convert_metafile(data: bytes) -> BytesIO | None:
    """Best-effort conversion of EMF/WMF to PNG via Pillow."""
    try:
        from PIL import Image

        img = Image.open(BytesIO(data))
        buf = BytesIO()
        img.save(buf, format="PNG")
        buf.seek(0)
        return buf
    except Exception:
        return None
