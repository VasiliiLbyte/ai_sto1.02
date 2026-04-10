"""Read DOCX files and extract structured content with tables and images."""

from __future__ import annotations

import base64
import logging
import re
from io import BytesIO
from pathlib import Path
from typing import Optional

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Emu
from lxml import etree

from models.document import (
    Alignment,
    DocumentContent,
    ImageData,
    ParagraphData,
    RunData,
    TableCellData,
    TableData,
)

logger = logging.getLogger(__name__)

EMU_PER_CM = 914400 / 2.54  # ~360000


def read_docx(filepath: str | Path) -> DocumentContent:
    """Parse a .docx file into a DocumentContent model."""
    filepath = Path(filepath)
    if not filepath.exists():
        raise FileNotFoundError(f"File not found: {filepath}")

    doc = Document(str(filepath))
    paragraphs = _extract_paragraphs(doc)
    tables = _extract_tables(doc)
    images = _extract_images(doc)

    return DocumentContent(
        filename=filepath.name,
        total_paragraphs=len(paragraphs),
        total_tables=len(tables),
        total_images=len(images),
        paragraphs=paragraphs,
        tables=tables,
        images=images,
    )


def _alignment_from_docx(val) -> Optional[Alignment]:
    if val is None:
        return None
    mapping = {0: Alignment.LEFT, 1: Alignment.CENTER, 2: Alignment.RIGHT, 3: Alignment.JUSTIFY}
    if isinstance(val, int):
        return mapping.get(val)
    name = str(val).upper()
    for a in Alignment:
        if a.value in name:
            return a
    return None


def _emu_to_cm(emu) -> Optional[float]:
    if emu is None:
        return None
    return round(emu / 914400 * 2.54, 2)


def _extract_paragraphs(doc: Document) -> list[ParagraphData]:
    result: list[ParagraphData] = []
    for i, p in enumerate(doc.paragraphs):
        runs = []
        for r in p.runs:
            runs.append(
                RunData(
                    text=r.text,
                    font_name=r.font.name,
                    font_size_pt=round(r.font.size / 12700, 1) if r.font.size else None,
                    bold=bool(r.font.bold),
                    italic=bool(r.font.italic),
                    underline=bool(r.font.underline),
                )
            )

        pf = p.paragraph_format
        is_heading = p.style.name.startswith("Heading")
        heading_level = None
        if is_heading:
            try:
                heading_level = int(p.style.name.split()[-1])
            except ValueError:
                heading_level = 1

        result.append(
            ParagraphData(
                index=i,
                text=p.text,
                runs=runs,
                style_name=p.style.name,
                alignment=_alignment_from_docx(p.alignment),
                first_line_indent_cm=_emu_to_cm(pf.first_line_indent),
                left_indent_cm=_emu_to_cm(pf.left_indent),
                space_before_pt=_emu_to_cm(pf.space_before),
                space_after_pt=_emu_to_cm(pf.space_after),
                line_spacing=pf.line_spacing if isinstance(pf.line_spacing, (int, float)) else None,
                is_heading=is_heading,
                heading_level=heading_level,
            )
        )
    return result


def _extract_tables(doc: Document) -> list[TableData]:
    """Extract tables while tracking their position relative to paragraphs."""
    tables: list[TableData] = []
    body = doc.element.body

    para_counter = 0
    table_idx = 0
    for child in body:
        tag = etree.QName(child.tag).localname
        if tag == "p":
            para_counter += 1
        elif tag == "tbl":
            tbl_obj = _find_table_by_xml(doc, child)
            if tbl_obj is None:
                continue
            rows_data: list[list[TableCellData]] = []
            for row in tbl_obj.rows:
                row_cells: list[TableCellData] = []
                for cell in row.cells:
                    cell_paragraphs = []
                    for cp in cell.paragraphs:
                        cell_paragraphs.append(
                            ParagraphData(text=cp.text, runs=[
                                RunData(
                                    text=r.text,
                                    font_name=r.font.name,
                                    font_size_pt=round(r.font.size / 12700, 1) if r.font.size else None,
                                    bold=bool(r.font.bold),
                                    italic=bool(r.font.italic),
                                )
                                for r in cp.runs
                            ])
                        )
                    row_cells.append(
                        TableCellData(text=cell.text.strip(), paragraphs=cell_paragraphs)
                    )
                rows_data.append(row_cells)

            tables.append(
                TableData(
                    index=table_idx,
                    position_after_paragraph=para_counter,
                    rows=rows_data,
                    num_rows=len(tbl_obj.rows),
                    num_cols=len(tbl_obj.columns),
                )
            )
            table_idx += 1

    return tables


def _find_table_by_xml(doc: Document, tbl_element) -> Optional[object]:
    """Find the python-docx Table object matching a given XML element."""
    for table in doc.tables:
        if table._tbl is tbl_element:
            return table
    return None


def _extract_images(doc: Document) -> list[ImageData]:
    """Extract inline images from the document."""
    images: list[ImageData] = []
    img_idx = 0

    for i, p in enumerate(doc.paragraphs):
        for run in p.runs:
            drawing_els = run._element.findall(qn("w:drawing"))
            for drawing in drawing_els:
                blip_els = drawing.findall(".//" + qn("a:blip"))
                for blip in blip_els:
                    embed = blip.get(qn("r:embed"))
                    if not embed:
                        continue
                    rel = doc.part.rels.get(embed)
                    if rel is None:
                        continue
                    try:
                        img_bytes = rel.target_part.blob
                    except Exception:
                        continue

                    content_type = rel.target_part.content_type or "image/png"
                    ext = _ext_from_content_type(content_type)

                    extent = drawing.find(".//" + qn("wp:extent"))
                    w_cm = h_cm = None
                    if extent is not None:
                        cx = extent.get("cx")
                        cy = extent.get("cy")
                        if cx:
                            w_cm = round(int(cx) / 914400 * 2.54, 2)
                        if cy:
                            h_cm = round(int(cy) / 914400 * 2.54, 2)

                    img = ImageData(
                        index=img_idx,
                        position_after_paragraph=i,
                        content_type=content_type,
                        width_cm=w_cm,
                        height_cm=h_cm,
                        filename=f"image_{img_idx}{ext}",
                    )
                    img.set_bytes(img_bytes)
                    images.append(img)
                    img_idx += 1

    return images


def _ext_from_content_type(ct: str) -> str:
    mapping = {
        "image/png": ".png",
        "image/jpeg": ".jpg",
        "image/gif": ".gif",
        "image/x-wmf": ".wmf",
        "image/x-emf": ".emf",
        "image/tiff": ".tiff",
    }
    return mapping.get(ct, ".png")
